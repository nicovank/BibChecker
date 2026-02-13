__all__ = ["get_colors", "write_output", "write_multi_output"]

from docx.shared import RGBColor
import re


RESET  = "\033[0m"
def get_colors(doc):
    if doc:
        RED    = RGBColor(0xD1, 0x24, 0x2F)   
        ORANGE = RGBColor(0xBC, 0x4C, 0x00)   
        GREEN  = RGBColor(0x2E, 0x7D, 0x32)
        BLUE   = RGBColor(0x1F, 0x6F, 0xEB)   
        DIM    = RGBColor(0x88, 0x88, 0x88)
    else:
        RED    = "\033[31m"
        ORANGE = "\033[33m"  
        GREEN  = "\033[32m"
        BLUE   = "\033[34m"
        DIM    = "\033[2m"

    return [RED, ORANGE, BLUE, GREEN, DIM]

def clean_xml_text(text):
    if not text:
        return text

    # Remove NULL bytes
    text = text.replace('\x00', '')

    # Remove all control chars except \t \n \r
    text = re.sub(r'[\x01-\x08\x0B-\x0C\x0E-\x1F]', '', text)

    return text

def write_output(string, doc, color=None):
    if doc:
        p = doc.add_paragraph()
        string = clean_xml_text(string)
        run = p.add_run(f"{string}")
        if color:
            run.font.color.rgb = color
    else:
        if color is None:
            color = RESET
        print(f"{color}{string}{RESET}")


def write_multi_output(header, header_color, colored_words, doc):
    if doc:
        p = doc.add_paragraph()
        header = clean_xml_text(header)
        run = p.add_run(header)
        run.font.color.rgb = header_color
        for word, color in colored_words:
            word = clean_xml_text(word)
            run = p.add_run(word)
            if color is not None:
                run.font.color.rgb = color
            run = p.add_run(" ")
    else:
        print(f"{header_color}{header}{RESET} ", end="")
        for word, color in colored_words:
            if color is None:
                color = RESET
            print(f"{color}{word}{RESET} ", end="")
        print("\n")
